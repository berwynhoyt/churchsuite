#!/usr/bin/env python3
""" Fetch upcoming service plans from ChurchSuite and export into docx format """

import sys
import os
import pprint
import argparse
import logging
import textwrap
import re

from datetime import date, timedelta
from types import SimpleNamespace

import churchsuite
import pathvalidate

# For docx export
import docx
from docx.shared import RGBColor
from docx.shared import Mm, Pt
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml.shared import qn

__version__ = '1.0.0'

# Regex of pattern used to identify start of red-highlighted text in service plans
everyone_pattern = re.compile(r'(.* |^)((all|everyone|together|^people):)(.*)', re.IGNORECASE + re.DOTALL)
# Regex of pattern used to identify start on non-red text in service plans
leader_pattern = re.compile(r'(.* |^)((leader|minister|reader):)(.*)', re.IGNORECASE + re.DOTALL)

def add_paragraph(doc, words):
    red = RGBColor(255, 0, 0)
    color = None

    para = doc.add_paragraph()
    # Add words to paragraph a line at a time so we can mark text said by all as red
    lines = words.splitlines(keepends=True)
    prev_line = None
    for line in lines:
        original_line = line
        if line.strip() in [':', '.', '-']:
            line = line.replace(':', '').replace('.', '').replace('-', '')
        # Blank line (i.e. \n\n) reverts to normal text
        if line.strip() == '':
            color = None
        # Lines intended for all to recite switch to red text
        match = everyone_pattern.search(line)
        if match:
            para.add_run(match.group(1) + match.group(2)).bold = True
            line = match.group(4)
            color = red
        # Lines for the leader are in black text
        match = leader_pattern.search(line)
        if match:
            color = None
            para.add_run(match.group(1) + match.group(2)).bold = True
            line = match.group(4)
        addition = para.add_run(line)
        if color:
            addition.font.color.rgb = color
        if line.strip().endswith(':') and not prev_line:
            addition.bold = True
        prev_line = original_line

def set_language(doc, language):
    # Access the default run properties element
    styles_element = doc.styles.element
    rpr_default = styles_element.xpath('./w:docDefaults/w:rPrDefault/w:rPr')[0]
    # Access or create the w:lang element and set the language value
    lang_default = rpr_default.xpath('w:lang')[0]
    lang_default.set(qn('w:val'), language) # Example: set to German (Germany)


def item_sections(item):
    """ Return a dictionary of all the named sections of the given service plan item """
    sections = {}
    responses = getattr(item, 'question_responses')
    for q in responses or ():
        if not q:
            continue
        sections[q.name] = q.value.replace('\r\n', '\n')
        
    return sections

# code to add page number

def set_page_size(section, size):
    """ Set page size to "width,height" in mm or "A4" or "letter" """
    size = size.lower()
    if size != 'letter': # do nothing if it's letter because that's the docx default
        if size == 'a4': size = "210,297"
    width, height = size.split(',')
    section.page_height = Mm(int(height))
    section.page_width = Mm(int(width))
    section.left_margin = Mm(22)
    section.right_margin = Mm(22)
    section.top_margin = Mm(22)
    section.bottom_margin = Mm(22)

def add_page_number(section):
    from docx.oxml import OxmlElement, ns
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    # I don't know how this works. It was taken from: https://stackoverflow.com/questions/56658872/add-page-number-using-python-docx
    def create_element(name): return OxmlElement(name)
    def create_attribute(element, name, value): element.set(ns.qn(name), value)

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    paragraph = section.header.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = paragraph.add_run()
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def plan2docx(cs, plan, stream=None, quiet=False):
    """ Save service plan to MS Word .docx for clearer presentation and markup by the service leader.
        The output is also less noisy than the pdf plan exported by ChurchSuite.
        If stream of type io.BytesIO() is supplied, save to stream instead of to a filename matching the plan.
        Return filename matching the name of the service plan.
    """
    green = RGBColor(0, 153, 0)
    black = RGBColor(0, 0, 0)

    title = f"{plan.date} {plan.name}{' (draft)' if plan.status=='draft' else ''}"
    filename = pathvalidate.sanitize_filename(title) + '.docx'
    if not quiet:
        print(f"Creating {filename}")
    doc = docx.Document()
    set_language(doc, args.language)
    doc.styles['Normal'].font.size = Pt(args.fontsize)
    doc.styles['Heading 1'].font.size = Pt(args.fontsize + 3)
    doc.styles['Heading 2'].font.size = Pt(args.fontsize + 1)
    section = doc.sections[0]
    set_page_size(section, args.pagesize)
    add_page_number(section)
    # Calculate the position of the right margin (page width - left margin - right margin) for right-margin tabstop below
    right_margin = section.page_width - section.left_margin - section.right_margin

    doc.add_heading(title, level=0)
    items = cs.get(f'{churchsuite.api}/planning/plan_items', params={'plan_ids[]':plan.id})
    for item in items:
        logging.info(pprint.pformat(item))
        names = [f"{person.first_name} {person.last_name}" for person in item.people or []]
        heading = doc.add_heading(level=1)
        run = heading.add_run(item.name)
        # If it's a song we put it in green with song title
        if item.name.lower().strip() in ['song', 'psalm', 'hymn']:
            run.bold = False
            run.font.color.rgb = green
            if item.comment:
                run = heading.add_run(': ' + item.comment.strip())
                run.font.color.rgb = green

        # Add a new tab stop at the calculated end position, with RIGHT alignment
        tab_stops = heading.paragraph_format.tab_stops
        tab_stops.add_tab_stop(right_margin, alignment=WD_TAB_ALIGNMENT.RIGHT)
        if names:
            run = heading.add_run('\t(' + ', '.join(names) + ')')
            run.font.size = doc.styles['Normal'].font.size
            run.font.color.rgb = black
            run.bold = False

        sections = item_sections(item)
        for section, words in sections.items():
            if len(sections) > 1 and words.strip():
                doc.add_heading(section, level=2)
            if words:
                add_paragraph(doc, words)

    if stream:
        doc.save(stream)
        stream.seek(0)  # rewind so stream can be read back or sent out as a web response
    else:
        doc.save(filename)
    return filename

def plan2txt(cs, plan):
    """ Print service plan as txt. This is mainly for developer tinkering. """
    print(f"{plan.date} {plan.name} {' (draft)' if plan.status=='draft' else ''}:")
    items = cs.get(f'{churchsuite.api}/planning/plan_items', params={'plan_ids[]':plan.id})
    for item in items:
        names = [f"{person.first_name} {person.last_name}" for person in item.people or []]
        if names:
            item.name += ' (' + ', '.join(names) + ')'
        print(item.name)

        logging.info(pprint.pformat(item))
        for section, words in item_sections(item).items():
            print(textwrap.indent(f"*{section}*: {words}", '  '))

def get_serviceplans(cs):
    """ Get all published and draft plans for the next week and output them as Word .docx files for easier markup.
        These are less noisy than the default Churchsuite pdf plans.
        Use args.starts_after and args.starts_before to set dates of plans to get.
    """
    today = date.today().isoformat()
    starts_after, starts_before = args.starts_after, args.starts_before
    if not starts_after and not starts_before:
        starts_after = today
    if starts_after == 'today': starts_after = today
    if starts_before == 'today': starts_before = today
    kwargs = {}
    if starts_after:
        # make start date inclusive of that date
        kwargs['starts_after'] = (date.fromisoformat(starts_after) - timedelta(days=1)).isoformat()
    if starts_before:
        kwargs['starts_before'] = starts_before
    plans = []
    for status in ('published', 'draft'):
        plans += cs.get(f'{churchsuite.api}/planning/plans', status=status, **kwargs)
    return plans

# Set defaults that may be used instead of command-line parameters when this module is imported (e.g. by docexport_app.py)
# Note: type of default must match desired type as docexport_app.py uses that fact to cast incoming query parameters
args = SimpleNamespace(language='en_AU', pagesize='A4', fontsize=14, starts_after='', starts_before='')

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument('--starts-after', action='store', default='', help="Specify start date (YYYY-MM-DD) from which to download upcoming service plans (default today).")
    parser.add_argument('--starts-before', action='store', default='', help="Specify end date (YYYY-MM-DD) to download service plans up to.")
    parser.add_argument('--language', action='store', default=args.language, help="Set language for docx file.")
    parser.add_argument('--pagesize', action='store', default=args.pagesize, help='Set page size to "width,height" in mm or "A4" or "letter".')
    parser.add_argument('--fontsize', action='store', type=int, default=args.fontsize, help='Set normal fontsize on Pt. Headings are enlargements of this.')
    parser.add_argument('-v', '--verbose', action='count', default=0, help="Increase verbosity level (e.g., -vv).")
    parser.add_argument('--txt', action='store_true', help="Output text to terminal rather than to a docx file.")
    parser.add_argument('--raw', action='store', default=None, help="Send all json received from the server into the specified raw json file.")
    parser.add_argument('--version', action='store_true', help="Print version number of this script and exit.")
    args = parser.parse_args()
    if len(sys.argv) < 2:
        print(f"{sys.argv[0]} exports ChurchSuite service plans to a docx file. For help, run: {sys.argv[0]} -h")

    # Set logging level based on -v flag
    log_level = logging.WARNING - 10*args.verbose
    logging.basicConfig(level=log_level, format=f'%(levelname)s: %(message)s')

    if args.version:
        print(__version__)
        sys.exit()

    import config
    cs = churchsuite.Churchsuite(auth=(config.USER_CLIENT_ID, config.USER_CLIENT_SECRET), raw=args.raw)
    plans = get_serviceplans(cs)
    if not plans:
        sys.exit(f"There are no plans in ChurchSuite starting after ({args.starts_after if args.starts_after or args.starts_before else 'today'}) and before ({args.starts_before})")
    for plan in plans:
        if args.txt:
            plan2txt(cs, plan)
        else:
            plan2docx(cs, plan)

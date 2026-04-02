#!/usr/bin/env python3
""" Fetch the next service plan from ChurchSuite and output it as a word document """

import pprint
import argparse
import logging
import textwrap
import re

from datetime import date, timedelta

import churchsuite as cs
import secrets
import pathvalidate

import docx
from docx.shared import RGBColor
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml.shared import qn

# Regex of pattern used to identify start of red-highlighted text in service plans
everyone_pattern = re.compile(r'(.* |^)((all|everyone|together|^people):)(.*)', re.IGNORECASE + re.DOTALL)
# Regex of pattern used to identify start on non-red text in service plans
leader_pattern = re.compile(r'(.* |^)((leader|minister|reader):)(.*)', re.IGNORECASE + re.DOTALL)

def add_paragraph(doc, words):
    red = RGBColor(255, 0, 0)
    color = None

    para = doc.add_paragraph()
    # Add words to paragraph a line at a time so we can mark text said by all as red
    lines = words.splitlines(keepends=True)
    prev_line = None
    for line in lines:
        original_line = line
        if line.strip() in [':', '.', '-']:
            line = line.replace(':', '').replace('.', '').replace('-', '')
        # Blank line (i.e. \n\n) reverts to normal text
        if line.strip() == '':
            color = None
        # Lines intended for all to recite switch to red text
        match = everyone_pattern.search(line)
        if match:
            para.add_run(match.group(1) + match.group(2)).bold = True
            line = match.group(4)
            color = red
        # Lines for the leader are in black text
        match = leader_pattern.search(line)
        if match:
            color = None
            para.add_run(match.group(1) + match.group(2)).bold = True
            line = match.group(4)
        addition = para.add_run(line)
        if color:
            addition.font.color.rgb = color
        if line.strip().endswith(':') and not prev_line:
            addition.bold = True
        prev_line = original_line

def set_language(doc, language):
    # Access the default run properties element
    styles_element = doc.styles.element
    rpr_default = styles_element.xpath('./w:docDefaults/w:rPrDefault/w:rPr')[0]
    # Access or create the w:lang element and set the language value
    lang_default = rpr_default.xpath('w:lang')[0]
    lang_default.set(qn('w:val'), language) # Example: set to German (Germany)


def item_sections(item):
    """ Return a dictionary of all the named sections of the given service plan item """
    sections = {}
    responses = getattr(item, 'question_responses')
    if not responses and item.comment:
        sections['comment'] = item.comment.replace('\r\n', '\n')
        return sections

    for q in responses or ():
        if not q:
            continue
        sections[q.name] = q.value.replace('\r\n', '\n')
        
    return sections

def plan2docx(plan, quiet=False):
    """ Save service plan to MS Word .docx file for easier markup by the service leader.
        The output is also less noisy than the pdf plan exported by ChurchSuite.
    """
    title = f"{plan.date} {plan.name}{' (draft)' if plan.status=='draft' else ''}"
    filename = pathvalidate.sanitize_filename(title) + '.docx'
    if not quiet:
        print(f"Creating {filename}")
    doc = docx.Document()
    set_language(doc, args.language)
    # Calculate the position of the right margin (page width - left margin - right margin)
    sec = doc.sections[0]
    margin_end = sec.page_width - sec.left_margin - sec.right_margin

    doc.add_heading(title)
    items = db.get(cs.URL.plan_items, params={'plan_ids[]':plan.id})
    for item in items:
        names = [f"{person.first_name} {person.last_name}" for person in item.people or []]
        item_name = item.name
        if names:
            item_name += '\t(' + ', '.join(names) + ')'
        heading = doc.add_heading(item_name, level=1)
        # Add a new tab stop at the calculated end position, with RIGHT alignment
        tab_stops = heading.paragraph_format.tab_stops
        tab_stops.add_tab_stop(margin_end, alignment=WD_TAB_ALIGNMENT.RIGHT)

        logging.info(pprint.pformat(item))
        sections = item_sections(item)
        for section, words in sections.items():
            if len(sections) > 1:
                doc.add_heading(section, level=2)
            if words:
                add_paragraph(doc, words)

    doc.save(filename)
    return filename

def plan2txt(plan):
    """ Print service plan as txt. This is mainly for developer tinkering. """
    print(f"{plan.date} {plan.name} {' (draft)' if plan.status=='draft' else ''}:")
    items = db.get(cs.URL.plan_items, params={'plan_ids[]':plan.id})
    for item in items:
        names = [f"{person.first_name} {person.last_name}" for person in item.people or []]
        item_name = item.name
        if names:
            item_name += ' (' + ', '.join(names) + ')'
        print(item_name)

        logging.info(pprint.pformat(item))
        for section, words in item_sections(item).items():
            print(textwrap.indent(f"*{section}*: {words}", '  '))

def upcoming_services(db):
    """ Get all published and draft plans for the next week and output them as Word .docx files for easier markup.
        These are less noisy than the default Churchsuite 
    """
    today = date.today()
    plans = []
    for status in ('published', 'draft'):
        plans += db.get(cs.URL.plans, status=status, starts_after=str(today), starts_before=str(today + timedelta(days=7)))
    if not plans:
        sys.exit("There is no plan in ChurchSuite for the coming week")
    return plans

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument('-v', '--verbose', action='count', default=0, help="Increase verbosity level (e.g., -vv)")
    parser.add_argument('--txt', action='store_true', help="Output text to terminal rather than to a docx file")
    parser.add_argument('--raw', action='store', default=None, help="Send all json received from the server into the specified raw json file")
    parser.add_argument('--language', action='store', default='en-AU', help='Set language for docx file')
    args = parser.parse_args()
    # Set logging level based on -v flag
    log_level = logging.WARNING - 10*args.verbose
    logging.basicConfig(level=log_level, format=f'%(levelname)s: %(message)s')

    db = cs.Churchsuite(secrets.CLIENT_ID, secrets.CLIENT_SECRET, raw=args.raw)
    #db = cs.Churchsuite(secrets.CLIENT_ID_app, secrets.CLIENT_SECRET_app, redirect_url="https://stgilesgreenwich.churchsuite.com")

    for plan in upcoming_services(db):
        if args.txt:
            plan2txt(plan)
        else:
            plan2docx(plan)
